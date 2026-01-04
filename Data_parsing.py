

import os
import re
import json
import pdfplumber
import docx


# ------------------------------------------------
# 1. Read a .txt resume file and print full content
# ------------------------------------------------
resume_txt = open("resume.txt", "r", encoding="utf-8").read()
print("Q1 Resume Full Content:\n", resume_txt)


# ------------------------------------------------
# 2. Display first 300 characters as preview
# ------------------------------------------------
print("\nQ2 Resume Preview (300 chars):\n", resume_txt[:300])


# ------------------------------------------------
# 3. Convert text to lowercase and remove extra spaces
# ------------------------------------------------
clean_resume_txt = re.sub(r"\s+", " ", resume_txt.lower()).strip()
print("\nQ3 Cleaned Text:\n", clean_resume_txt)


# ------------------------------------------------
# 4. Detect file type (PDF / DOCX / TXT)
# ------------------------------------------------
def detect_file_type(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return "PDF"
    elif ext == ".docx":
        return "DOCX"
    elif ext == ".txt":
        return "TXT"
    else:
        return "Unsupported"

print("\nQ4 File Type:", detect_file_type("resume.txt"))


# ------------------------------------------------
# 5. Load resume and job description (TXT)
# ------------------------------------------------
job_txt = open("job.txt", "r", encoding="utf-8").read()
print("\nQ5 Resume & Job Description Loaded Successfully")


# ------------------------------------------------
# 6. Extract text from PDF resume
# ------------------------------------------------
pdf_text = ""
with pdfplumber.open("resume.pdf") as pdf:
    for page in pdf.pages:
        if page.extract_text():
            pdf_text += page.extract_text()

print("\nQ6 PDF Resume Text Extracted")


# ------------------------------------------------
# 7. Extract paragraphs from DOCX resume
# ------------------------------------------------
doc = docx.Document("resume.docx")
docx_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
print("\nQ7 DOCX Resume Text Extracted")


# ------------------------------------------------
# 8. Unified function to load PDF, DOCX, TXT
# ------------------------------------------------
def load_file(path):
    if not os.path.exists(path):
        raise FileNotFoundError("Invalid file path")

    if path.endswith(".txt"):
        return open(path, "r", encoding="utf-8").read()

    elif path.endswith(".pdf"):
        text = ""
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                if page.extract_text():
                    text += page.extract_text()
        return text

    elif path.endswith(".docx"):
        doc = docx.Document(path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    else:
        raise ValueError("Unsupported file type")


# ------------------------------------------------
# 9. Display resume & job preview (500 chars)
# ------------------------------------------------
resume_data = load_file("resume.txt")
job_data = load_file("job.txt")

print("\nQ9 Resume Preview:\n", resume_data[:500])
print("\nQ9 Job Preview:\n", job_data[:500])


# ------------------------------------------------
# 10. Display raw and cleaned text
# ------------------------------------------------
clean_resume = re.sub(r"\s+", " ", resume_data.lower()).strip()
print("\nQ10 RAW TEXT:\n", resume_data)
print("\nQ10 CLEANED TEXT:\n", clean_resume)


# ------------------------------------------------
# 11. Simulate file upload using file paths
# ------------------------------------------------
resume_path = "resume.pdf"
job_path = "job.txt"

resume_uploaded = load_file(resume_path)
job_uploaded = load_file(job_path)
print("\nQ11 File Upload Simulation Successful")


# ------------------------------------------------
# 12. Handle error cases
# ------------------------------------------------
try:
    load_file("resume.jpg")
except Exception as e:
    print("\nQ12 Error Handled:", e)


# ------------------------------------------------
# 13. Store parsed data in dictionary
# ------------------------------------------------
parsed_data = {
    "resume": resume_uploaded,
    "job_description": job_uploaded
}
print("\nQ13 Parsed Data Stored in Dictionary")


# ------------------------------------------------
# 14. Save cleaned text into JSON
# ------------------------------------------------
cleaned_output = {
    "resume_cleaned": re.sub(r"\s+", " ", resume_uploaded.lower()).strip(),
    "job_cleaned": re.sub(r"\s+", " ", job_uploaded.lower()).strip()
}

with open("parsed_output.json", "w", encoding="utf-8") as f:
    json.dump(cleaned_output, f, indent=4)

print("\nQ14 Cleaned Data Saved to JSON")


# ------------------------------------------------
# 15. Parse multiple resumes (different formats)
# ------------------------------------------------
resume_files = ["resume.txt", "resume.pdf", "resume.docx"]
resume_list = []

for r in resume_files:
    resume_list.append(re.sub(r"\s+", " ", load_file(r).lower()).strip())

print("\nQ15 Multiple Resumes Parsed")


# ------------------------------------------------
# 16. Universal parser (detect + extract + clean + preview)
# ------------------------------------------------
def universal_parser(path):
    raw = load_file(path)
    cleaned = re.sub(r"\s+", " ", raw.lower()).strip()
    preview = raw[:300]
    return raw, cleaned, preview


# ------------------------------------------------
# 17. Parse multiple resumes with one job description
# ------------------------------------------------
for r in resume_files:
    _, _, prev = universal_parser(r)
    print(f"\nQ17 Preview of {r}:\n", prev)


# ------------------------------------------------
# 18. Preview formatting module
# ------------------------------------------------
def format_preview(text, length=300):
    return text[:length]


# ------------------------------------------------
# 19. Replicate complete Milestone-1 output
# ------------------------------------------------
print("\nQ19 Milestone-1 Output Generated Successfully")


# ------------------------------------------------
# 20. End-to-end execution completed
# ------------------------------------------------
print("\nQ20 End-to-End Data Parsing Script Executed Successfully")
